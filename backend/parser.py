import logging
import re
from pathlib import Path
from typing import Iterable

import fitz
from docx import Document


ALLOWED_EXTENSIONS = {".pdf", ".docx"}
EMPTY_JOB_DESCRIPTION_MESSAGE = "Enter a Job Description to calculate Job Match."
NO_RECOGNIZABLE_SKILLS_MESSAGE = "No recognizable skills found in Job Description."

logger = logging.getLogger(__name__)

TECH_SKILLS = [
    "python",
    "java",
    "javascript",
    "typescript",
    "c",
    "c++",
    "c#",
    "go",
    "ruby",
    "php",
    "swift",
    "kotlin",
    "sql",
    "mysql",
    "postgresql",
    "sqlite",
    "mongodb",
    "redis",
    "html",
    "css",
    "tailwind css",
    "bootstrap",
    "react",
    "redux",
    "next.js",
    "vue",
    "angular",
    "node.js",
    "nodejs",
    "express",
    "django",
    "flask",
    "fastapi",
    "spring boot",
    "laravel",
    "rest api",
    "graphql",
    "api",
    "docker",
    "kubernetes",
    "jenkins",
    "git",
    "github",
    "linux",
    "aws",
    "azure",
    "google cloud",
    "gcp",
    "firebase",
    "terraform",
    "ansible",
    "ci/cd",
    "machine learning",
    "deep learning",
    "nlp",
    "natural language processing",
    "data science",
    "data analysis",
    "data engineering",
    "pandas",
    "numpy",
    "scikit-learn",
    "tensorflow",
    "pytorch",
    "opencv",
    "socketio",
    "power bi",
    "tableau",
    "excel",
    "figma",
    "ui/ux",
    "agile",
    "scrum",
    "jira",
    "project management",
    "leadership",
    "communication",
    "problem solving",
]

DISPLAY_NAMES = {
    "api": "API",
    "aws": "AWS",
    "c": "C",
    "c#": "C#",
    "c++": "C++",
    "ci/cd": "CI/CD",
    "css": "CSS",
    "fastapi": "FastAPI",
    "gcp": "GCP",
    "github": "GitHub",
    "google cloud": "Google Cloud",
    "graphql": "GraphQL",
    "html": "HTML",
    "javascript": "JavaScript",
    "mongodb": "MongoDB",
    "mysql": "MySQL",
    "next.js": "Next.js",
    "nlp": "NLP",
    "node.js": "Node.js",
    "nodejs": "NodeJS",
    "opencv": "OpenCV",
    "php": "PHP",
    "postgresql": "PostgreSQL",
    "power bi": "Power BI",
    "rest api": "REST API",
    "scikit-learn": "Scikit-learn",
    "sql": "SQL",
    "sqlite": "SQLite",
    "socketio": "SocketIO",
    "tailwind css": "Tailwind CSS",
    "typescript": "TypeScript",
    "ui/ux": "UI/UX",
}

SECTION_HEADERS = {
    "career_objective": [
        "career objective",
        "objective",
        "summary",
        "professional summary",
    ],
    "skills": [
        "skills",
        "technical skills",
        "core skills",
        "key skills",
        "technologies",
        "tools and technologies",
    ],
    "education": [
        "education",
        "academic background",
        "academics",
        "qualification",
        "qualifications",
    ],
    "experience": [
        "experience",
        "internship/experience",
        "internship experience",
        "work experience",
        "professional experience",
        "employment",
        "career history",
        "internship",
        "internships",
    ],
    "projects": [
        "projects",
        "project experience",
        "personal projects",
        "academic projects",
    ],
    "certifications": [
        "certifications",
        "certificates",
        "licenses",
        "courses",
        "achievements",
        "awards",
    ],
    "personal_details": [
        "personal details",
        "personal detail",
        "personal information",
        "personal profile",
    ],
}

EDUCATION_PATTERN = re.compile(
    r"\b(bachelor|master|b\.?tech|m\.?tech|b\.?e\.?|m\.?e\.?|bsc|msc|mba|"
    r"ph\.?d|diploma|degree|university|college|school|cgpa|gpa)\b",
    re.IGNORECASE,
)

CERTIFICATION_PATTERN = re.compile(
    r"\b(certified|certification|certificate|aws certified|azure certified|"
    r"google cloud|microsoft certified|oracle certified|coursera|udemy)\b",
    re.IGNORECASE,
)

URL_PATTERN = re.compile(r"https?://[^\s)>\]]+", re.IGNORECASE)
YEAR_PATTERN = re.compile(r"^(?:19|20)\d{2}$")
YEAR_ANYWHERE_PATTERN = re.compile(r"\b(?:19|20)\d{2}\b")
PROJECT_ACTION_VERBS = {
    "added",
    "automated",
    "built",
    "collaborated",
    "created",
    "deployed",
    "designed",
    "developed",
    "engineered",
    "enhanced",
    "implemented",
    "improved",
    "integrated",
    "led",
    "managed",
    "optimized",
    "reduced",
    "tested",
    "trained",
    "used",
}
PROJECT_LINK_LABEL_PATTERN = re.compile(
    r"^(live|demo|deployed|website|app|github|git hub|repo|repository)\s*:",
    re.IGNORECASE,
)
PROJECT_TECH_LABEL_PATTERN = re.compile(
    r"^(tech stack|technologies|technology|tools|stack|built with)\s*:",
    re.IGNORECASE,
)
PROJECT_TITLE_SEPARATOR_PATTERN = re.compile(r"\s*(?:\u2014|\u2013|\ufffd)\s*|\s(?:-|\?)\s")
PROJECT_DURATION_PATTERN = re.compile(
    r"^(?:(?:jan|feb|mar|apr|may|jun|jul|aug|sep|sept|oct|nov|dec)[a-z]*\s+)?"
    r"(?:19|20)\d{2}\s*(?:-|–|—|\?|to)\s*"
    r"(?:(?:jan|feb|mar|apr|may|jun|jul|aug|sep|sept|oct|nov|dec)[a-z]*\s+)?"
    r"(?:19|20)\d{2}$",
    re.IGNORECASE,
)
PROJECT_STATUS_PATTERN = re.compile(
    r"^(?:(?:19|20)\d{2}\s+)?\(?\s*(in development|under development|in progress|ongoing|current|working)\s*\)?(?:\s+(?:19|20)\d{2})?$",
    re.IGNORECASE,
)
CERTIFICATION_DATE_PATTERN = re.compile(
    r"\b(?:jan|feb|mar|apr|may|jun|jul|aug|sep|sept|oct|nov|dec)[a-z]*\s+(?:19|20)\d{2}\b",
    re.IGNORECASE,
)
PERSONAL_DETAIL_PATTERN = re.compile(
    r"\b(father|mother|contact|phone|email|address|date of birth|dob|gender|marital|nationality|language|candidate name)\b",
    re.IGNORECASE,
)


def extract_text_from_file(file_path: str | Path) -> str:
    """Read text from a PDF or DOCX resume."""
    path = Path(file_path)
    extension = path.suffix.lower()

    if extension == ".pdf":
        return clean_text(extract_text_from_pdf(path))
    if extension == ".docx":
        return clean_text(extract_text_from_docx(path))

    raise ValueError("Only PDF and DOCX files are supported.")


def extract_text_from_pdf(file_path: Path) -> str:
    text_parts: list[str] = []
    with fitz.open(file_path) as document:
        for page in document:
            text_parts.append(page.get_text("text"))
    return "\n".join(text_parts)


def extract_text_from_docx(file_path: Path) -> str:
    document = Document(file_path)
    text_parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]

    for table in document.tables:
        for row in table.rows:
            row_text = " ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)

    return "\n".join(text_parts)


def parse_resume(raw_text: str) -> dict:
    """Parse common resume fields from raw text."""
    text = clean_text(raw_text)
    sections = extract_section_blocks(text)
    skills = extract_skills(text)

    parsed = {
        "name": extract_name(text),
        "email": extract_email(text),
        "phone": extract_phone(text),
        "skills": skills,
        "education": extract_education(text, sections.get("education", "")),
        "experience": extract_experience(sections.get("experience", "")),
        "projects": extract_projects(sections.get("projects", "")),
        "certifications": extract_certifications(text, sections.get("certifications", "")),
        "raw_text": text,
    }
    parsed["ats_score"] = calculate_ats_score(parsed)
    return parsed


def calculate_match_score(candidate_skills: Iterable[str], job_description: str) -> dict:
    """Compare JD skills with candidate skills using the shared skill dictionary."""
    extracted_resume_skills = sorted(list(candidate_skills), key=str.lower)
    extracted_jd_skills = extract_skills(job_description)

    if not job_description.strip():
        logger.info("extracted_resume_skills=%s", extracted_resume_skills)
        logger.info("extracted_jd_skills=%s", [])
        logger.info("matched_skills=%s", [])
        logger.info("missing_skills=%s", [])
        return {
            "job_skills": [],
            "matched_skills": [],
            "missing_skills": [],
            "match_score": None,
            "score": None,
            "message": EMPTY_JOB_DESCRIPTION_MESSAGE,
        }

    if not extracted_jd_skills:
        logger.info("extracted_resume_skills=%s", extracted_resume_skills)
        logger.info("extracted_jd_skills=%s", [])
        logger.info("matched_skills=%s", [])
        logger.info("missing_skills=%s", [])
        return {
            "job_skills": [],
            "matched_skills": [],
            "missing_skills": [],
            "match_score": None,
            "score": None,
            "message": NO_RECOGNIZABLE_SKILLS_MESSAGE,
        }

    resume_skill_map = {skill.lower(): skill for skill in extracted_resume_skills}
    job_skill_map = {skill.lower(): skill for skill in extracted_jd_skills}

    matched_keys = sorted(set(resume_skill_map) & set(job_skill_map))
    missing_keys = sorted(set(job_skill_map) - set(resume_skill_map))

    matched_skills = [job_skill_map[key] for key in matched_keys]
    missing_skills = [job_skill_map[key] for key in missing_keys]
    match_score = round((len(matched_skills) / len(extracted_jd_skills)) * 100, 2)

    logger.info("extracted_resume_skills=%s", extracted_resume_skills)
    logger.info("extracted_jd_skills=%s", sorted(extracted_jd_skills, key=str.lower))
    logger.info("matched_skills=%s", sorted(matched_skills, key=str.lower))
    logger.info("missing_skills=%s", sorted(missing_skills, key=str.lower))

    return {
        "job_skills": sorted(extracted_jd_skills, key=str.lower),
        "matched_skills": sorted(matched_skills, key=str.lower),
        "missing_skills": sorted(missing_skills, key=str.lower),
        "match_score": match_score,
        "score": match_score,
        "message": "",
    }


def match_job_description(resume_skills: Iterable[str], job_description: str) -> dict:
    return calculate_match_score(resume_skills, job_description)


def calculate_ats_score(parsed_resume: dict) -> int:
    """Simple ATS-style score based on profile completeness and skill density."""
    score = 0
    score += 10 if parsed_resume.get("name") and parsed_resume["name"] != "Unknown Candidate" else 0
    score += 15 if parsed_resume.get("email") else 0
    score += 10 if parsed_resume.get("phone") else 0
    score += min(len(parsed_resume.get("skills", [])) * 4, 25)
    score += 10 if parsed_resume.get("education") else 0
    score += 15 if parsed_resume.get("experience") else 0
    score += 10 if parsed_resume.get("projects") else 0
    score += 5 if parsed_resume.get("certifications") else 0
    return min(score, 100)


def clean_text(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_email(text: str) -> str:
    match = re.search(r"[\w.+-]+@[\w-]+\.[\w.-]+", text)
    return match.group(0).strip(".") if match else ""


def extract_phone(text: str) -> str:
    phone_patterns = [
        r"(?:\+?\d{1,3}[\s.-]?)?(?:\(?\d{3}\)?[\s.-]?)\d{3}[\s.-]?\d{4}",
        r"(?:\+?\d{1,3}[\s.-]?)?\d{5}[\s.-]?\d{5}",
    ]
    for pattern in phone_patterns:
        match = re.search(pattern, text)
        if match:
            return re.sub(r"\s+", " ", match.group(0)).strip()
    return ""


def extract_name(text: str) -> str:
    for line in get_clean_lines(text)[:12]:
        if "@" in line or re.search(r"\d{5,}", line):
            continue

        candidate = re.sub(r"[^A-Za-z\s.'-]", " ", line)
        candidate = re.sub(r"\s+", " ", candidate).strip(" .-")
        words = candidate.split()

        if not 2 <= len(words) <= 5:
            continue
        if any(word.lower() in {"resume", "curriculum", "vitae", "email", "phone"} for word in words):
            continue
        if sum(word[:1].isupper() for word in words) >= 2:
            return candidate

    return "Unknown Candidate"


def extract_skills(text: str) -> list[str]:
    lowered = text.lower()
    found: set[str] = set()

    for skill in TECH_SKILLS:
        pattern = rf"(?<![a-z0-9+#]){re.escape(skill)}(?![a-z0-9+#])"
        if re.search(pattern, lowered):
            found.add(format_skill_name(skill))

    return sorted(found, key=str.lower)


def extract_section_blocks(text: str) -> dict[str, str]:
    """Return full text blocks for known resume sections."""
    sections: dict[str, list[str]] = {section: [] for section in SECTION_HEADERS}
    current_section: str | None = None

    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            if current_section and sections[current_section]:
                sections[current_section].append("")
            continue

        matched_section = match_section_heading(line)
        if matched_section:
            current_section = matched_section
            continue

        if current_section:
            sections[current_section].append(raw_line.rstrip())

    return {
        section: clean_text("\n".join(lines))
        for section, lines in sections.items()
    }


def match_section_heading(line: str) -> str | None:
    normalized = normalize_heading(line)
    heading_to_section = {
        heading: section
        for section, headings in SECTION_HEADERS.items()
        for heading in headings
    }

    for heading in sorted(heading_to_section, key=len, reverse=True):
        if normalized == heading or normalized.startswith(f"{heading} "):
            return heading_to_section[heading]

    return None


def extract_sections(text: str) -> dict[str, list[str]]:
    """Backward-compatible line view of the section blocks."""
    return {
        section: get_clean_lines(block)
        for section, block in extract_section_blocks(text).items()
    }


def extract_education(text: str, section_block: str) -> list[dict[str, str] | str]:
    section_lines = get_clean_lines(section_block)
    if section_lines:
        parsed_rows = parse_education_rows(section_lines)
        return parsed_rows if parsed_rows else section_lines[:10]

    return [line for line in get_clean_lines(text) if EDUCATION_PATTERN.search(line)][:8]


def parse_education_rows(lines: list[str]) -> list[dict[str, str]]:
    content_lines = [
        line
        for line in lines
        if not is_education_header_line(line)
    ]

    parsed_from_rows = [
        row
        for line in content_lines
        if (row := parse_education_row(line))
    ]
    if parsed_from_rows:
        return parsed_from_rows

    grouped_rows = group_education_cells(content_lines)
    return [
        row
        for line in grouped_rows
        if (row := parse_education_row(line))
    ]


def group_education_cells(lines: list[str]) -> list[str]:
    grouped: list[str] = []
    current: list[str] = []

    for line in lines:
        if starts_education_level(line):
            if current:
                grouped.append(" ".join(current))
            current = [line]
        elif current:
            current.append(line)

    if current:
        grouped.append(" ".join(current))

    return grouped


def parse_education_row(line: str) -> dict[str, str] | None:
    levels = ["High School", "Intermediate", "Graduation", "Post Graduation", "Diploma"]
    cleaned = re.sub(r"\s+", " ", line).strip()
    lowered = cleaned.lower()

    matched_level = next((level for level in levels if lowered.startswith(level.lower())), "")
    if not matched_level:
        return None

    remainder = cleaned[len(matched_level):].strip()
    parts = remainder.split()
    if len(parts) < 3:
        return None

    result = parts[-1]
    year = parts[-2]
    board_university = " ".join(parts[:-2])

    return {
        "level": matched_level,
        "board_university": board_university,
        "year": year,
        "result": result,
    }


def starts_education_level(line: str) -> bool:
    return bool(re.match(r"^(graduation|intermediate|high school|post graduation|diploma)\b", line, re.IGNORECASE))


def is_education_header_line(line: str) -> bool:
    lowered = line.lower()
    return (
        "education" in lowered
        and "board" in lowered
        and "year" in lowered
        and "result" in lowered
    ) or lowered in {"education", "board/university", "board", "university", "year", "result"}


def extract_experience(section_block: str) -> list[str]:
    return get_clean_lines(section_block)[:15]


def extract_projects(section_block: str | list[str]) -> list[dict[str, object]]:
    """Group project section lines into structured project objects."""
    section_lines = (
        get_clean_lines(section_block)
        if isinstance(section_block, str)
        else section_block
    )
    projects: list[dict[str, object]] = []
    current_project: dict[str, object] | None = None

    for raw_line in section_lines:
        line = clean_project_line(raw_line)
        if not line:
            continue

        if is_project_title(line):
            if current_project and has_project_content(current_project):
                projects.append(finalize_project(current_project))
            current_project = create_project(line)
            continue

        if current_project is None:
            current_project = create_project("Project")

        add_line_to_project(current_project, line)

    if current_project and has_project_content(current_project):
        projects.append(finalize_project(current_project))

    return projects


def create_project(title_line: str) -> dict[str, object]:
    title, year = parse_project_title(title_line)
    return {
        "title": title or "Project",
        "year": year,
        "duration": "",
        "live_url": "",
        "github_url": "",
        "tech_stack": [],
        "description": [],
    }


def parse_project_title(line: str) -> tuple[str, str]:
    cleaned = clean_project_line(line)
    year_match = YEAR_ANYWHERE_PATTERN.search(cleaned)
    year = year_match.group(0) if year_match else ""

    title = YEAR_ANYWHERE_PATTERN.sub("", cleaned)
    title = re.sub(r"\s+\?\s+", " — ", title)
    title = re.sub(r"\s+", " ", title).strip(" -:|")
    return title, year


def add_line_to_project(project: dict[str, object], line: str) -> None:
    if is_project_status(line):
        year_match = YEAR_ANYWHERE_PATTERN.search(line)
        if year_match and not project.get("year"):
            project["year"] = year_match.group(0)
        return

    if YEAR_PATTERN.fullmatch(line) and not project.get("year"):
        project["year"] = line
        return

    if is_project_duration(line):
        project["duration"] = normalize_project_duration(line)
        return

    urls = extract_urls(line)
    if urls:
        assign_project_url(project, line, urls[0])
        return

    if is_technology_line(line):
        extend_unique(project["tech_stack"], parse_technologies(line))
        return

    append_unique(project["description"], clean_project_description(line))


def assign_project_url(project: dict[str, object], line: str, url: str) -> None:
    lowered = line.lower()
    if "github" in lowered or "git hub" in lowered or "github.com" in lowered:
        project["github_url"] = url
        return

    if PROJECT_LINK_LABEL_PATTERN.search(line) or not project.get("live_url"):
        project["live_url"] = url
        return

    append_unique(project["description"], line)


def is_project_title(line: str) -> bool:
    if has_url(line) or YEAR_PATTERN.fullmatch(line) or is_project_duration(line):
        return False
    if is_project_status(line):
        return False
    if PROJECT_LINK_LABEL_PATTERN.search(line) or PROJECT_TECH_LABEL_PATTERN.search(line):
        return False
    if is_technology_line(line) or starts_with_action_verb(line):
        return False
    if line.endswith("."):
        return False

    if is_year_dash_project_title(line):
        return True

    if len(line.split()) > 16:
        return False
    if PROJECT_TITLE_SEPARATOR_PATTERN.search(line):
        return True
    if ":" in line:
        prefix = line.split(":", 1)[0].strip()
        return 0 < len(prefix.split()) <= 8
    return looks_like_title_case(line)


def is_year_dash_project_title(line: str) -> bool:
    has_dash = bool(PROJECT_TITLE_SEPARATOR_PATTERN.search(line))
    has_year = bool(YEAR_ANYWHERE_PATTERN.search(line))
    return has_dash and has_year


def is_technology_line(line: str) -> bool:
    if has_url(line) or YEAR_PATTERN.fullmatch(line) or is_project_duration(line) or starts_with_action_verb(line):
        return False
    if PROJECT_TECH_LABEL_PATTERN.search(line) or "|" in line:
        return True

    extracted_skills = extract_skills(line)
    is_short_line = len(line.split()) <= 14 and not line.endswith(".")
    has_separators = "," in line or " / " in line or "\u2022" in line or "\u00b7" in line
    return is_short_line and (len(extracted_skills) >= 3 or (len(extracted_skills) >= 2 and has_separators))


def parse_technologies(line: str) -> list[str]:
    cleaned = PROJECT_TECH_LABEL_PATTERN.sub("", line).strip()

    if "|" in cleaned or "," in cleaned or " / " in cleaned or "\u2022" in cleaned or "\u00b7" in cleaned:
        parts = re.split(r"\s*(?:\||,|\u2022|\u00b7| / )\s*", cleaned)
        technologies = [normalize_technology(part) for part in parts]
        return [technology for technology in technologies if technology]

    return extract_skills_in_text_order(cleaned)


def normalize_technology(value: str) -> str:
    cleaned = clean_project_line(value).strip(" .;")
    if not cleaned:
        return ""

    lowered = cleaned.lower()
    return format_skill_name(lowered) if lowered in TECH_SKILLS else cleaned


def extract_skills_in_text_order(text: str) -> list[str]:
    lowered = text.lower()
    matches: list[tuple[int, str]] = []

    for skill in TECH_SKILLS:
        pattern = rf"(?<![a-z0-9+#]){re.escape(skill)}(?![a-z0-9+#])"
        match = re.search(pattern, lowered)
        if match:
            matches.append((match.start(), format_skill_name(skill)))

    return [skill for _, skill in sorted(matches, key=lambda item: item[0])]


def looks_like_title_case(line: str) -> bool:
    words = re.findall(r"[A-Za-z0-9][A-Za-z0-9+#./'-]*", line)
    if not 1 <= len(words) <= 10:
        return False
    if line.lower() in {"live demo", "github", "repository", "source code"}:
        return False

    title_like_words = sum(
        1
        for word in words
        if word[:1].isupper() or word.isupper() or any(character.isdigit() for character in word)
    )
    return title_like_words >= min(2, len(words))


def starts_with_action_verb(line: str) -> bool:
    first_word_match = re.match(r"[A-Za-z]+", line.strip())
    return bool(first_word_match and first_word_match.group(0).lower() in PROJECT_ACTION_VERBS)


def extract_urls(line: str) -> list[str]:
    return [match.group(0).rstrip(".,;") for match in URL_PATTERN.finditer(line)]


def has_url(line: str) -> bool:
    return bool(URL_PATTERN.search(line))


def is_project_duration(line: str) -> bool:
    normalized = re.sub(r"\s+", " ", line).strip()
    return bool(PROJECT_DURATION_PATTERN.fullmatch(normalized))


def is_project_status(line: str) -> bool:
    normalized = re.sub(r"\s+", " ", line).strip()
    return bool(PROJECT_STATUS_PATTERN.fullmatch(normalized))


def normalize_project_duration(line: str) -> str:
    return re.sub(r"\s*(?:\?|-|–|—|to)\s*", " – ", line, count=1, flags=re.IGNORECASE).strip()


def clean_project_line(line: str) -> str:
    cleaned = re.sub(r"\s+", " ", line).strip()
    return cleaned.strip(" \t\r\n-*•●▪▫")


def clean_project_description(line: str) -> str:
    return line.strip(" \t\r\n-*•●▪▫?").strip()


def append_unique(items: object, value: str) -> None:
    if not isinstance(items, list):
        return
    key = value.lower()
    if value and all(existing.lower() != key for existing in items if isinstance(existing, str)):
        items.append(value)


def extend_unique(items: object, values: Iterable[str]) -> None:
    for value in values:
        append_unique(items, value)


def has_project_content(project: dict[str, object]) -> bool:
    return any(
        [
            project.get("title") and project.get("title") != "Project",
            project.get("year"),
            project.get("duration"),
            project.get("live_url"),
            project.get("github_url"),
            project.get("tech_stack"),
            project.get("description"),
        ]
    )


def finalize_project(project: dict[str, object]) -> dict[str, object]:
    project["tech_stack"] = unique_items(project.get("tech_stack", []))
    project["description"] = unique_items(project.get("description", []))
    return project


def extract_certifications(text: str, section_block: str) -> list[dict[str, str] | str]:
    section_lines = get_clean_lines(section_block)
    if section_lines:
        parsed_certifications = parse_certifications(section_lines)
        return parsed_certifications if parsed_certifications else section_lines[:10]

    return [line for line in get_clean_lines(text) if CERTIFICATION_PATTERN.search(line)][:8]


def parse_certifications(lines: list[str]) -> list[dict[str, str]]:
    certifications: list[dict[str, str]] = []

    for line in lines:
        if should_skip_certification_line(line):
            continue

        date_match = CERTIFICATION_DATE_PATTERN.search(line)
        if date_match:
            date = date_match.group(0)
            name = CERTIFICATION_DATE_PATTERN.sub("", line).strip(" -:|")
            if name:
                certifications.append({"name": name, "date": date})
            elif certifications and not certifications[-1].get("date"):
                certifications[-1]["date"] = date
            continue

        if certifications and not certifications[-1].get("date") and CERTIFICATION_DATE_PATTERN.fullmatch(line):
            certifications[-1]["date"] = line
            continue

        certifications.append({"name": line, "date": ""})

    return certifications[:10]


def should_skip_certification_line(line: str) -> bool:
    normalized = normalize_heading(line)
    if normalized in {"certifications", "certificates"}:
        return True
    if match_section_heading(line) == "personal_details":
        return True
    return bool(PERSONAL_DETAIL_PATTERN.search(line))


def get_clean_lines(text: str) -> list[str]:
    lines = []
    for raw_line in text.splitlines():
        line = raw_line.strip()
        line = re.sub(r"\s+", " ", line)
        line = line.strip(" \t\r\n-*•●▪▫")
        if line:
            lines.append(line)
    return lines


def normalize_heading(line: str) -> str:
    line = line.strip().lower()
    line = re.sub(r"^[^\w]+|[^\w]+$", "", line)
    line = re.sub(r"\s+", " ", line)
    return line


def unique_items(items: Iterable[str]) -> list[str]:
    seen = set()
    result = []
    for item in items:
        cleaned = item.strip()
        key = cleaned.lower()
        if cleaned and key not in seen:
            seen.add(key)
            result.append(cleaned)
    return result


def format_skill_name(skill: str) -> str:
    if skill in DISPLAY_NAMES:
        return DISPLAY_NAMES[skill]
    return " ".join(part.capitalize() for part in skill.split())
